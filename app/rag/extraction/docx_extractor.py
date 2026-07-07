from pathlib import Path
from typing import List

from docx import Document as DocxDocument

from app.core.config.logging_config import get_logger
from app.rag.extraction.base import BaseExtractor, ExtractedPage

log = get_logger(__name__)

# DOCX has no native "page" concept in the XML (pagination is a rendering-time
# thing). We approximate pages by grouping paragraphs, which keeps citations
# meaningful without needing a full layout engine.
PARAGRAPHS_PER_PSEUDO_PAGE = 25


class DOCXExtractor(BaseExtractor):
    @property
    def supported_extensions(self) -> List[str]:
        return [".docx"]

    def extract(self, file_path: Path) -> List[ExtractedPage]:
        try:
            doc = DocxDocument(str(file_path))
        except Exception as exc:
            log.error("docx_extraction_failed", file=str(file_path), error=str(exc))
            raise ValueError(f"Failed to extract DOCX '{file_path.name}': {exc}") from exc

        # Include tables too — enterprise docs (policies, handbooks) often
        # carry key facts in tables.
        blocks: List[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    blocks.append(row_text)

        pages: List[ExtractedPage] = []
        for i in range(0, len(blocks), PARAGRAPHS_PER_PSEUDO_PAGE):
            group = blocks[i : i + PARAGRAPHS_PER_PSEUDO_PAGE]
            page_number = (i // PARAGRAPHS_PER_PSEUDO_PAGE) + 1
            pages.append(ExtractedPage(page_number=page_number, text="\n".join(group)))

        if not pages:
            log.warning("docx_no_extractable_text", file=str(file_path))
        return pages
